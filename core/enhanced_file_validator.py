"""
Enhanced File Validation Service

This module provides comprehensive file validation functionality to address POQ-002:
- Magic number validation using python-magic
- Empty and corrupt file detection for DOCX and PDF formats
- Content validation beyond simple extension checking
- Integration with existing validation infrastructure
"""

from __future__ import annotations

import io
import tempfile
from pathlib import Path
from typing import Any, NamedTuple

import docx
import fitz  # PyMuPDF

from legal_portal.core.logging_config import setup_logging


logger = setup_logging("enhanced_file_validator")

# Optional import for python-magic (fallback to extension-based validation if not available)
try:
    import magic

    MAGIC_AVAILABLE = True
    logger.info(
        "python-magic library available - enhanced magic number validation enabled"
    )
except ImportError:
    MAGIC_AVAILABLE = False
    logger.warning(
        "python-magic library not available - falling back to extension-based validation"
    )


class ValidationResult(NamedTuple):
    """Result of file validation"""

    is_valid: bool
    detected_type: str | None
    file_size: int
    issues: list[str]
    warnings: list[str]


class ValidationError(Exception):
    """Raised when file validation encounters an error"""


class EnhancedFileValidator:
    """
    Enhanced file validation service with magic number detection and corruption checking.

    Addresses POQ-002 by providing:
    - Magic number validation to verify file type matches extension
    - Enhanced empty file detection (beyond zero bytes)
    - Corrupt file detection for DOCX and PDF formats
    - Comprehensive logging for debugging validation issues
    """

    # Supported file types and their magic number signatures
    SUPPORTED_TYPES = {
        "application/pdf": [".pdf"],
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": [
            ".docx"
        ],
        "application/msword": [".doc"],
        "text/plain": [".txt"],
        "message/rfc822": [".eml"],
        "image/png": [".png"],
        "image/jpeg": [".jpg", ".jpeg"],
    }

    # Minimum file sizes for different types (in bytes)
    MIN_FILE_SIZES = {
        "application/pdf": 100,  # Minimum for a valid PDF header
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": 500,  # Minimum for ZIP structure
        "application/msword": 512,  # Minimum for OLE structure
        "text/plain": 1,  # At least one character
        "message/rfc822": 50,  # Basic email headers
        "image/png": 67,  # Minimum for PNG header + IHDR chunk
        "image/jpeg": 10,  # Minimum for JPEG header
    }

    def __init__(self):
        """Initialize the enhanced file validator"""
        self.magic_mime = None
        self.magic_available = MAGIC_AVAILABLE

        if self.magic_available:
            try:
                self.magic_mime = magic.Magic(mime=True)
                logger.info(
                    "Enhanced file validator initialized with magic number support"
                )
            except Exception as e:
                logger.warning(f"Failed to initialize python-magic: {e}")
                self.magic_available = False

        if not self.magic_available:
            logger.info(
                "Enhanced file validator initialized with extension-based validation only"
            )

    def validate_file(self, file_data: bytes, filename: str) -> ValidationResult:
        """
        Comprehensive file validation including magic numbers and corruption detection.

        Args:
            file_data: Raw file content as bytes
            filename: Original filename for extension checking

        Returns:
            ValidationResult with validation status and details
        """
        logger.debug(
            "Starting enhanced file validation",
            extra={
                "filename": filename,
                "file_size": len(file_data),
                "validation_type": "enhanced_with_magic"
                if MAGIC_AVAILABLE
                else "extension_based",
            },
        )

        issues = []
        warnings = []
        detected_type = None

        try:
            # Basic file size check
            file_size = len(file_data)
            if file_size == 0:
                issues.append("File is empty (0 bytes)")
                logger.warning(
                    "Empty file detected",
                    extra={"filename": filename, "validation_issue": "zero_bytes"},
                )
                return ValidationResult(False, None, file_size, issues, warnings)

            # Extract file extension
            file_extension = Path(filename).suffix.lower()
            if not file_extension:
                issues.append("File has no extension")
                logger.warning(
                    "File without extension",
                    extra={"filename": filename, "validation_issue": "no_extension"},
                )

            # Magic number validation if available
            if self.magic_available and self.magic_mime:
                detected_type = self._validate_magic_numbers(
                    file_data, file_extension, issues, warnings
                )
            else:
                detected_type = self._validate_extension_only(file_extension, issues)

            # Type-specific validation
            if detected_type:
                self._validate_file_content(
                    file_data, detected_type, filename, issues, warnings
                )
                self._validate_minimum_size(file_data, detected_type, issues)

            # Determine overall validation result
            is_valid = len(issues) == 0

            if is_valid:
                logger.info(
                    "File validation passed",
                    extra={
                        "filename": filename,
                        "detected_type": detected_type,
                        "file_size": file_size,
                        "warnings_count": len(warnings),
                    },
                )
            else:
                logger.error(
                    "File validation failed",
                    extra={
                        "filename": filename,
                        "issues": issues,
                        "warnings": warnings,
                        "file_size": file_size,
                    },
                )

            return ValidationResult(
                is_valid, detected_type, file_size, issues, warnings
            )

        except Exception as e:
            error_msg = f"Validation error: {e!s}"
            issues.append(error_msg)
            logger.error(
                "Exception during file validation",
                extra={
                    "filename": filename,
                    "error": str(e),
                    "validation_issue": "exception",
                },
            )
            return ValidationResult(False, None, len(file_data), issues, warnings)

    def _validate_magic_numbers(
        self,
        file_data: bytes,
        file_extension: str,
        issues: list[str],
        warnings: list[str],
    ) -> str | None:
        """Validate file using magic numbers"""
        try:
            # Use first 2048 bytes for accurate detection (as recommended by python-magic docs)
            detection_bytes = file_data[:2048] if len(file_data) > 2048 else file_data
            detected_mime = self.magic_mime.from_buffer(detection_bytes)

            logger.debug(
                "Magic number detection completed",
                extra={
                    "detected_mime": detected_mime,
                    "file_extension": file_extension,
                    "bytes_analyzed": len(detection_bytes),
                },
            )

            # Check if detected type is supported
            if detected_mime not in self.SUPPORTED_TYPES:
                issues.append(f"Unsupported file type detected: {detected_mime}")
                return None

            # Check if extension matches detected type
            expected_extensions = self.SUPPORTED_TYPES[detected_mime]
            if file_extension not in expected_extensions:
                if file_extension in [
                    ext for exts in self.SUPPORTED_TYPES.values() for ext in exts
                ]:
                    # Extension is supported but doesn't match content
                    issues.append(
                        f"File extension '{file_extension}' does not match detected type '{detected_mime}'. "
                        f"Expected extensions: {', '.join(expected_extensions)}"
                    )
                else:
                    # Extension is not supported at all
                    warnings.append(
                        f"Unusual file extension '{file_extension}' for type '{detected_mime}'"
                    )

            return detected_mime

        except Exception as e:
            warnings.append(f"Magic number detection failed: {e!s}")
            logger.warning(
                "Magic number detection failed",
                extra={"error": str(e), "fallback": "extension_based"},
            )
            return self._validate_extension_only(file_extension, issues)

    def _validate_extension_only(
        self, file_extension: str, issues: list[str]
    ) -> str | None:
        """Fallback validation using file extension only"""
        # Map extensions to MIME types
        extension_map = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".txt": "text/plain",
            ".eml": "message/rfc822",
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
        }

        detected_type = extension_map.get(file_extension)
        if not detected_type:
            issues.append(f"Unsupported file extension: {file_extension}")
            return None

        logger.debug(
            "Extension-based validation completed",
            extra={"file_extension": file_extension, "mapped_type": detected_type},
        )

        return detected_type

    def _validate_minimum_size(
        self, file_data: bytes, detected_type: str, issues: list[str]
    ) -> None:
        """Validate file meets minimum size requirements"""
        file_size = len(file_data)
        min_size = self.MIN_FILE_SIZES.get(detected_type, 1)

        if file_size < min_size:
            issues.append(
                f"File size ({file_size} bytes) below minimum for {detected_type} ({min_size} bytes)"
            )
            logger.warning(
                "File below minimum size threshold",
                extra={
                    "detected_type": detected_type,
                    "file_size": file_size,
                    "minimum_size": min_size,
                    "validation_issue": "undersized_file",
                },
            )

    def _validate_file_content(
        self,
        file_data: bytes,
        detected_type: str,
        filename: str,
        issues: list[str],
        warnings: list[str],
    ) -> None:
        """Validate file content for corruption and emptiness"""
        if (
            detected_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            self._validate_docx_content(file_data, filename, issues, warnings)
        elif detected_type == "application/pdf":
            self._validate_pdf_content(file_data, filename, issues, warnings)
        elif detected_type == "text/plain":
            self._validate_text_content(file_data, filename, issues, warnings)
        elif detected_type == "image/png":
            self._validate_png_content(file_data, filename, issues, warnings)
        elif detected_type == "image/jpeg":
            self._validate_jpg_content(file_data, filename, issues, warnings)
        elif detected_type == "application/msword":
            self._validate_doc_content(file_data, filename, issues, warnings)

    def _validate_docx_content(
        self, file_data: bytes, filename: str, issues: list[str], warnings: list[str]
    ) -> None:
        """Validate DOCX file content and structure"""
        try:
            # Create temporary file for python-docx to read
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                temp_file.write(file_data)
                temp_file.flush()

                try:
                    # Attempt to open with python-docx
                    document = docx.Document(temp_file.name)

                    # Check if document has any readable content
                    paragraph_texts = [
                        para.text.strip() for para in document.paragraphs
                    ]
                    non_empty_paragraphs = [text for text in paragraph_texts if text]

                    if not non_empty_paragraphs:
                        # Check tables for content
                        table_content = []
                        for table in document.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    cell_text = cell.text.strip()
                                    if cell_text:
                                        table_content.append(cell_text)

                        if not table_content:
                            warnings.append(
                                "DOCX file appears to be empty (no readable text content)"
                            )
                            logger.warning(
                                "Empty DOCX content detected",
                                extra={
                                    "filename": filename,
                                    "validation_issue": "empty_docx_content",
                                    "paragraphs_count": len(document.paragraphs),
                                    "tables_count": len(document.tables),
                                },
                            )

                    logger.debug(
                        "DOCX content validation completed",
                        extra={
                            "filename": filename,
                            "paragraphs_count": len(document.paragraphs),
                            "non_empty_paragraphs": len(non_empty_paragraphs),
                            "tables_count": len(document.tables),
                        },
                    )

                finally:
                    # Clean up temporary file
                    Path(temp_file.name).unlink(missing_ok=True)

        except Exception as e:
            issues.append(f"DOCX file appears to be corrupt: {e!s}")
            logger.error(
                "DOCX corruption detected",
                extra={
                    "filename": filename,
                    "error": str(e),
                    "validation_issue": "docx_corruption",
                },
            )

    def _validate_pdf_content(
        self, file_data: bytes, filename: str, issues: list[str], warnings: list[str]
    ) -> None:
        """Validate PDF file content and structure"""
        try:
            # Use PyMuPDF to validate PDF structure
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(file_data)
                temp_file.flush()

                try:
                    with fitz.open(temp_file.name) as doc:
                        page_count = len(doc)

                        if page_count == 0:
                            warnings.append("PDF file has no pages")
                            logger.warning(
                                "Empty PDF detected",
                                extra={
                                    "filename": filename,
                                    "validation_issue": "empty_pdf",
                                    "page_count": page_count,
                                },
                            )
                        else:
                            # Check if PDF has any extractable text
                            total_text_length = 0
                            for page in doc:
                                text = page.get_text().strip()
                                total_text_length += len(text)

                            if total_text_length == 0:
                                warnings.append(
                                    "PDF file contains no extractable text (may be image-only)"
                                )
                                logger.info(
                                    "PDF with no extractable text",
                                    extra={
                                        "filename": filename,
                                        "page_count": page_count,
                                        "validation_issue": "pdf_no_text",
                                    },
                                )

                        logger.debug(
                            "PDF content validation completed",
                            extra={
                                "filename": filename,
                                "page_count": page_count,
                                "total_text_length": total_text_length,
                            },
                        )

                finally:
                    # Clean up temporary file
                    Path(temp_file.name).unlink(missing_ok=True)

        except Exception as e:
            issues.append(f"PDF file appears to be corrupt: {e!s}")
            logger.error(
                "PDF corruption detected",
                extra={
                    "filename": filename,
                    "error": str(e),
                    "validation_issue": "pdf_corruption",
                },
            )

    def _validate_text_content(
        self, file_data: bytes, filename: str, issues: list[str], warnings: list[str]
    ) -> None:
        """Validate text file content"""
        try:
            # Try to decode as UTF-8
            try:
                text_content = file_data.decode("utf-8").strip()
            except UnicodeDecodeError:
                # Try other common encodings
                for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                    try:
                        text_content = file_data.decode(encoding).strip()
                        warnings.append(f"Text file decoded using {encoding} encoding")
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    issues.append("Text file contains invalid character encoding")
                    return

            if not text_content:
                warnings.append("Text file appears to be empty")
                logger.warning(
                    "Empty text file detected",
                    extra={
                        "filename": filename,
                        "file_size": len(file_data),
                        "validation_issue": "empty_text_content",
                    },
                )

            logger.debug(
                "Text content validation completed",
                extra={
                    "filename": filename,
                    "content_length": len(text_content),
                    "file_size": len(file_data),
                },
            )

        except Exception as e:
            issues.append(f"Text file validation failed: {e!s}")
            logger.error(
                "Text file validation error",
                extra={
                    "filename": filename,
                    "error": str(e),
                    "validation_issue": "text_validation_error",
                },
            )

    def _validate_png_content(
        self, file_data: bytes, filename: str, issues: list[str], warnings: list[str]
    ) -> None:
        """Validate PNG file content and structure with magic number check"""
        try:
            # Check PNG magic number: 89 50 4E 47 0D 0A 1A 0A
            png_signature = b"\x89\x50\x4E\x47\x0D\x0A\x1A\x0A"
            if not file_data.startswith(png_signature):
                issues.append("Invalid PNG magic number signature")
                logger.error(
                    "PNG magic number validation failed",
                    extra={"filename": filename, "validation_issue": "invalid_png_signature"}
                )
                return

            # Use PIL to validate PNG structure
            from PIL import Image
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
                temp_file.write(file_data)
                temp_file.flush()

                try:
                    with Image.open(temp_file.name) as img:
                        # Verify it's actually a PNG
                        if img.format != "PNG":
                            issues.append(f"File claims to be PNG but detected as {img.format}")
                            return
                        
                        # Check for basic validity
                        width, height = img.size
                        if width == 0 or height == 0:
                            issues.append("PNG file has invalid dimensions")
                        
                        logger.debug(
                            "PNG content validation completed",
                            extra={
                                "filename": filename,
                                "dimensions": f"{width}x{height}",
                                "mode": img.mode,
                            },
                        )

                finally:
                    Path(temp_file.name).unlink(missing_ok=True)

        except Exception as e:
            issues.append(f"PNG file appears to be corrupt: {e!s}")
            logger.error(
                "PNG corruption detected",
                extra={
                    "filename": filename,
                    "error": str(e),
                    "validation_issue": "png_corruption",
                },
            )

    def _validate_jpg_content(
        self, file_data: bytes, filename: str, issues: list[str], warnings: list[str]
    ) -> None:
        """Validate JPG file content and structure with magic number check"""
        try:
            # Check JPG magic number: FF D8 FF
            if not file_data.startswith(b"\xFF\xD8\xFF"):
                issues.append("Invalid JPEG magic number signature")
                logger.error(
                    "JPEG magic number validation failed",
                    extra={"filename": filename, "validation_issue": "invalid_jpeg_signature"}
                )
                return

            # Use PIL to validate JPEG structure
            from PIL import Image
            with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as temp_file:
                temp_file.write(file_data)
                temp_file.flush()

                try:
                    with Image.open(temp_file.name) as img:
                        # Verify it's actually a JPEG
                        if img.format != "JPEG":
                            issues.append(f"File claims to be JPEG but detected as {img.format}")
                            return
                        
                        # Check for basic validity
                        width, height = img.size
                        if width == 0 or height == 0:
                            issues.append("JPEG file has invalid dimensions")
                        
                        logger.debug(
                            "JPEG content validation completed",
                            extra={
                                "filename": filename,
                                "dimensions": f"{width}x{height}",
                                "mode": img.mode,
                            },
                        )

                finally:
                    Path(temp_file.name).unlink(missing_ok=True)

        except Exception as e:
            issues.append(f"JPEG file appears to be corrupt: {e!s}")
            logger.error(
                "JPEG corruption detected",
                extra={
                    "filename": filename,
                    "error": str(e),
                    "validation_issue": "jpeg_corruption",
                },
            )

    def _validate_doc_content(
        self, file_data: bytes, filename: str, issues: list[str], warnings: list[str]
    ) -> None:
        """Validate legacy DOC file content and structure with magic number check"""
        try:
            # Check DOC magic number: D0 CF 11 E0 A1 B1 1A E1 (OLE compound document)
            ole_signature = b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"
            if not file_data.startswith(ole_signature):
                issues.append("Invalid DOC magic number signature (not an OLE compound document)")
                logger.error(
                    "DOC magic number validation failed",
                    extra={"filename": filename, "validation_issue": "invalid_doc_signature"}
                )
                return

            # Try to validate with oletools if available
            try:
                from oletools import olefile
                with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
                    temp_file.write(file_data)
                    temp_file.flush()

                    try:
                        if olefile.isOleFile(temp_file.name):
                            ole = olefile.OleFileIO(temp_file.name)
                            try:
                                # Check for Word document streams
                                if ole.exists("WordDocument"):
                                    logger.debug(
                                        "Valid legacy DOC file detected",
                                        extra={"filename": filename, "streams": ole.listdir()}
                                    )
                                else:
                                    warnings.append("OLE file may not be a valid Word document")
                            finally:
                                ole.close()
                        else:
                            issues.append("File is not a valid OLE compound document")
                    finally:
                        Path(temp_file.name).unlink(missing_ok=True)

            except ImportError:
                # oletools not available, just verify the magic number was correct
                logger.debug(f"oletools not available, basic magic number validation passed for {filename}")
                
        except Exception as e:
            issues.append(f"DOC file appears to be corrupt: {e!s}")
            logger.error(
                "DOC corruption detected",
                extra={
                    "filename": filename,
                    "error": str(e),
                    "validation_issue": "doc_corruption",
                },
            )


# Convenience functions for integration with existing code
def validate_uploaded_file(file_data: bytes, filename: str) -> ValidationResult:
    """
    Convenience function to validate an uploaded file.

    Args:
        file_data: Raw file content as bytes
        filename: Original filename

    Returns:
        ValidationResult with validation status and details
    """
    validator = EnhancedFileValidator()
    return validator.validate_file(file_data, filename)


def is_file_valid(file_data: bytes, filename: str) -> bool:
    """
    Simple boolean check for file validity.

    Args:
        file_data: Raw file content as bytes
        filename: Original filename

    Returns:
        True if file passes all validation checks
    """
    result = validate_uploaded_file(file_data, filename)
    return result.is_valid
