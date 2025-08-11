from __future__ import annotations

import io
import mimetypes
import os

import docx

from legal_portal.core.data_models import (
    DocumentType,
    FileMetadata,
    FileType,
    ProcessedDocument,
)
from legal_portal.core.logging_config import get_module_logger


logger = get_module_logger(__name__)


async def process_docx(
    file_path: str, document_type: DocumentType, original_filename: str
) -> ProcessedDocument:
    """
    Processes a DOCX file by extracting its content from a given path.
    """
    logger.debug(f"Processing DOCX: {original_filename}")

    try:
        with open(file_path, "rb") as f:
            document = docx.Document(io.BytesIO(f.read()))
            full_text = "\n".join([para.text for para in document.paragraphs])
    except Exception as e:
        logger.info(f"Could not process {original_filename} with python-docx: {e}")
        full_text = f"Could not extract content from {original_filename}."

    content_type, _ = mimetypes.guess_type(file_path)
    file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

    # Create proper FileMetadata object with required fields
    file_metadata = FileMetadata(filename=original_filename, size=file_size)

    logger.info(
        f"✅ Fixed: Created FileMetadata for {original_filename}, size: {file_size}"
    )

    return ProcessedDocument(
        file_name=original_filename,
        content=full_text,
        document_type=document_type,
        file_type=FileType.DOCX,
        metadata=file_metadata,
    )
