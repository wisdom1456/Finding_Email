"""Document management endpoints."""

import logging
from datetime import datetime
from typing import List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from fastapi.responses import JSONResponse
from pydantic import BaseModel

from legal_portal.api.dependencies import get_current_user, get_supabase_client, get_user_supabase_client
from legal_portal.core.data_models import DocumentStatus, DocumentType
from legal_portal.core.document_processor import DocumentProcessor, ValidationError
from legal_portal.services.file_processors.eml_processor import process_eml
from legal_portal.utils.google_vision_client import GoogleVisionClient
from legal_portal.utils.security import sanitize_text_for_db

router = APIRouter()
logger = logging.getLogger(__name__)


class BulkDeleteRequest(BaseModel):
    """Request model for bulk delete operation."""

    document_ids: List[str]


class BulkDeleteResponse(BaseModel):
    """Response model for bulk delete operation."""

    deleted_count: int
    failed_ids: List[str]
    errors: List[str]


class BulkExtractRequest(BaseModel):
    """Request model for bulk extraction operation."""

    case_id: str


class BulkExtractResponse(BaseModel):
    """Response model for bulk extraction operation."""

    extracted_count: int
    failed_count: int
    errors: List[str]


class VerifyDocumentRequest(BaseModel):
    """Request model for verifying/correcting document text."""

    manual_text: Optional[str] = None
    is_verified: bool = True
    is_flagged_as_junk: bool = False


class DocumentResponse(BaseModel):
    """Response model for a document."""

    id: str
    case_id: str
    file_name: str
    file_type: str
    file_size: int
    storage_path: Optional[str] = None
    status: DocumentStatus
    created_at: datetime
    updated_at: datetime
    is_verified: bool = False
    is_flagged_as_junk: bool = False
    extracted_text: Optional[str] = None
    manual_text: Optional[str] = None
    extraction_method: Optional[str] = None
    extraction_quality: Optional[str] = None
    extraction_error: Optional[str] = None
    metadata: Optional[dict] = None


@router.post("/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    case_id: str = Form(...),
    file: UploadFile = File(...),
    is_intake_form: bool = Form(False),
    extract_immediately: bool = Form(True),
    user=Depends(get_current_user),
    user_supabase=Depends(get_user_supabase_client),
    service_supabase=Depends(get_supabase_client),
):
    """Upload a document for a case with unified validation and compression.

    Args:
    ----
        case_id: ID of the case this document belongs to
        file: File to upload
        is_intake_form: Whether this is an intake form
        extract_immediately: Whether to extract text immediately after upload (default: True)
        user: Current authenticated user
        user_supabase: User-scoped Supabase client (for RLS)
        service_supabase: Service-scoped Supabase client (bypasses RLS)

    Returns:
    -------
        Created document metadata

    Raises:
    ------
        400: Validation error (size, type, content, security)
        404: Case not found
        500: Server error

    """
    import os
    import tempfile

    from legal_portal.services.file_processors.pdf_processor import process_pdf

    try:
        logger.debug(
            f"Upload document: user={user['id']}, case={case_id}, "
            f"file={file.filename}, type={file.content_type}"
        )

        # Verify case ownership (use user client for RLS)
        logger.debug("Verifying case ownership...")
        case_response = (
            user_supabase.table("cases").select("id").eq("id", case_id).eq("user_id", user["id"]).execute()
        )
        logger.debug(f"Case found: {bool(case_response.data)}")

        if not case_response.data:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Case not found")

        # Fetch user profile for blacklist
        profile_response = service_supabase.table("profiles").select("ai_preferences").eq("id", user["id"]).execute()
        blacklist = []
        if profile_response.data and profile_response.data[0].get("ai_preferences"):
            blacklist = profile_response.data[0]["ai_preferences"].get("blacklisted_documents", [])

        # Read file content
        file_content = await file.read()
        logger.debug(f"File size: {len(file_content)} bytes")

        # Use unified processor for validation, compression, and upload
        processor = DocumentProcessor()

        try:
            doc_record = await processor.process_and_upload(
                file_content=file_content,
                filename=file.filename,
                user_id=user["id"],
                case_id=case_id,
                supabase_client=service_supabase,
                is_intake_form=is_intake_form,
                content_type=file.content_type,
                blacklist=blacklist,
            )
        except ValidationError as e:
            # Return structured validation error
            error_response = {
                "code": e.error_code,
                "detail": str(e),
                "file_name": file.filename,
                "file_size_mb": e.file_size_mb,
            }
            logger.warning(f"Validation error: {e.error_code} - {str(e)}")
            return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content=error_response)

        # Create document record in database (use user client for RLS)
        logger.debug("Creating document record in database...")
        doc_response = user_supabase.table("documents").insert(doc_record).execute()

        if not doc_response.data:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to create document record"
            )

        created_doc = doc_response.data[0]
        document_id = created_doc["id"]
        logger.info(f"Document uploaded successfully: {document_id}")

        # Immediate text extraction (non-blocking approach - extract in background)
        if extract_immediately:
            try:
                file_type = doc_record.get("file_type", file.content_type)
                file_name = doc_record.get("file_name", file.filename)

                extracted_text = ""
                extraction_method = ""
                extraction_quality = "high"
                ocr_provider = None
                extraction_error = None
                page_count = None

                if file_type in ["application/pdf", "pdf"] or file_name.lower().endswith(".pdf"):
                    # Write to temp file for PDF processing
                    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                        tmp.write(file_content)
                        tmp_path = tmp.name

                    try:
                        result = await process_pdf(
                            file_path=tmp_path,
                            document_type=DocumentType.CASE_DOCUMENT,
                            original_filename=file_name,
                        )
                        extracted_text = result.content
                        extraction_method = result.extraction_method or "unknown"
                        extraction_quality = result.extraction_quality or "high"
                        ocr_provider = result.ocr_provider
                        extraction_error = result.extraction_error
                        page_count = result.page_count
                    finally:
                        os.unlink(tmp_path)

                elif file_type in ["text/plain", "txt"] or file_name.lower().endswith(".txt"):
                    # Plain text file
                    try:
                        extracted_text = file_content.decode("utf-8", errors="replace")
                        extraction_method = "direct_text"
                        extraction_quality = "high"
                    except Exception as e:
                        extraction_error = f"Failed to decode text: {e}"
                        extraction_method = "failed"
                        extraction_quality = "low"

                elif file_type in ["image/png", "image/jpeg", "image/jpg"] or file_name.lower().endswith((".png", ".jpg", ".jpeg")):
                    # Image file - use Google Vision OCR (with GPT-4o fallback)
                    try:
                        google_client = GoogleVisionClient.get_instance()
                        if google_client.is_available:
                            try:
                                import asyncio

                                from starlette.concurrency import run_in_threadpool

                                def do_google_ocr():
                                    return google_client.extract_text_from_image(file_content)

                                vision_text = await asyncio.wait_for(
                                    run_in_threadpool(do_google_ocr),
                                    timeout=30.0,
                                )

                                if vision_text and vision_text.strip():
                                    extracted_text = vision_text
                                    extraction_method = "Google Cloud Vision"
                                    extraction_quality = "high"
                                    ocr_provider = "google_vision"
                                    logger.info(f"Successfully extracted text from {file_name} using Google Vision")
                                else:
                                    raise ValueError("Google Vision returned empty text")
                            except Exception as google_err:
                                logger.warning(f"{file_name}: Google Vision failed ({google_err}). Falling back to GPT-4o Vision.")
                                raise  # Fall through to GPT-4o fallback
                        else:
                            raise ValueError("Google Vision client not available")
                    except Exception:
                        # Fallback to GPT-4o Vision
                        try:
                            import asyncio
                            import base64

                            from starlette.concurrency import run_in_threadpool

                            from legal_portal.utils.openai_client import OpenAIClient

                            logger.info(f"Using GPT-4o Vision for {file_name}")
                            openai_client = OpenAIClient()
                            client = openai_client.client

                            # Determine MIME type
                            if file_type in ["image/jpeg", "image/jpg"] or file_name.lower().endswith((".jpg", ".jpeg")):
                                mime_type = "image/jpeg"
                            else:
                                mime_type = "image/png"

                            base64_image = base64.b64encode(file_content).decode("utf-8")
                            prompt = (
                                f"Extract ALL text from this legal document image. "
                                f"Filename: {file_name}. "
                                "This is a scanned document that needs OCR text extraction. "
                                "Maintain the logical structure and layout. "
                                "If there are tables, preserve the row/column relationship. "
                                "Provide the text verbatim including all numbers, dates, and names. Do not summarize."
                            )

                            def gpt4o_ocr():
                                return client.chat.completions.create(
                                    model="gpt-4o",
                                    messages=[{
                                        "role": "user",
                                        "content": [
                                            {"type": "text", "text": prompt},
                                            {
                                                "type": "image_url",
                                                "image_url": {
                                                    "url": f"data:{mime_type};base64,{base64_image}",
                                                    "detail": "high"
                                                }
                                            },
                                        ]
                                    }],
                                    max_tokens=4000,
                                    temperature=0.0,
                                )

                            response = await asyncio.wait_for(
                                run_in_threadpool(gpt4o_ocr),
                                timeout=60.0,
                            )

                            extracted_text = response.choices[0].message.content
                            extraction_method = "GPT-4o Vision"
                            extraction_quality = "high"
                            ocr_provider = "openai"
                            logger.info(f"Successfully extracted text from {file_name} using GPT-4o Vision")
                        except Exception as ocr_err:
                            extraction_error = f"Image OCR failed: {str(ocr_err)}"
                            extraction_method = "failed"
                            extraction_quality = "low"
                            logger.error(f"Image extraction error for {file_name}: {ocr_err}")

                elif file_type in ["message/rfc822", "eml"] or file_name.lower().endswith(".eml"):
                    # Email file (.eml)
                    try:
                        import asyncio
                        import tempfile

                        from starlette.concurrency import run_in_threadpool

                        # Write to temp file for eml processing
                        with tempfile.NamedTemporaryFile(suffix=".eml", delete=False) as tmp:
                            tmp.write(file_content)
                            tmp_path = tmp.name

                        try:
                            # Process eml file asynchronously
                            result = await process_eml(
                                file_path=tmp_path,
                                document_type=DocumentType.CORRESPONDENCE,
                                original_filename=file_name,
                            )

                            extracted_text = result.content
                            extraction_method = result.extraction_method or "email_parser"
                            extraction_quality = result.extraction_quality or "high"
                            extraction_error = result.extraction_error

                            logger.info(f"Successfully extracted email content from {file_name}: {len(extracted_text)} chars")
                        finally:
                            # Clean up temp file
                            import os
                            try:
                                os.unlink(tmp_path)
                            except Exception:
                                pass
                    except Exception as e:
                        extraction_error = f"Email extraction failed: {str(e)}"
                        extraction_method = "failed"
                        extraction_quality = "low"
                        logger.error(f"Email extraction error for {file_name}: {e}")

                else:
                    # Other file types - mark as needing extraction
                    extraction_method = "pending"
                    extraction_quality = "unknown"

                # Sanitize extracted text to remove NULL characters that PostgreSQL can't store
                extracted_text = sanitize_text_for_db(extracted_text)

                # Update document with extraction results
                update_data = {
                    "extracted_text": extracted_text if extracted_text else None,
                    "extraction_method": extraction_method,
                    "extraction_quality": extraction_quality,
                    "ocr_provider": ocr_provider,
                    "extraction_error": extraction_error,
                    "page_count": page_count,
                    "extracted_at": datetime.utcnow().isoformat() if extracted_text else None,
                    "updated_at": datetime.utcnow().isoformat(),
                    "status": DocumentStatus.READY if extracted_text else DocumentStatus.EXTRACTION_FAILED,
                }

                user_supabase.table("documents").update(update_data).eq("id", document_id).execute()
                logger.info(
                    f"Extraction complete for {document_id}: method={extraction_method}, "
                    f"quality={extraction_quality}, chars={len(extracted_text)}"
                )

            except Exception as extract_err:
                # Log but don't fail the upload if extraction fails
                logger.warning(f"Immediate extraction failed for {document_id}: {extract_err}")
                # Update with error status
                user_supabase.table("documents").update(
                    {
                        "extraction_method": "failed",
                        "extraction_quality": "low",
                        "extraction_error": str(extract_err),
                        "status": DocumentStatus.EXTRACTION_FAILED,
                        "updated_at": datetime.utcnow().isoformat(),
                    }
                ).eq("id", document_id).execute()

        return created_doc

    except HTTPException:
        raise
    except ValidationError as e:
        # Catch any validation errors that slipped through
        error_response = {
            "code": e.error_code,
            "detail": str(e),
            "file_name": file.filename,
            "file_size_mb": e.file_size_mb,
        }
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content=error_response)
    except Exception as e:
        logger.error(f"Error in upload_document: {type(e).__name__} - {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Error uploading document: {str(e)}"
        ) from e


@router.get("/case/{case_id}", response_model=List[DocumentResponse])
async def list_documents_for_case(
    case_id: str, user=Depends(get_current_user), supabase=Depends(get_user_supabase_client)
):
    """List all documents for a specific case.

    Args:
    ----
        case_id: Case ID
        user: Current authenticated user
        supabase: Supabase client

    Returns:
    -------
        List of documents

    """
    try:
        # Verify case ownership
        case_response = (
            supabase.table("cases").select("id").eq("id", case_id).eq("user_id", user["id"]).execute()
        )

        if not case_response.data:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Case not found")

        # Get documents
        response = (
            supabase.table("documents")
            .select("*")
            .eq("case_id", case_id)
            .order("created_at", desc=False)
            .execute()
        )

        return response.data
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Error fetching documents: {str(e)}"
        ) from e


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: str, user=Depends(get_current_user), supabase=Depends(get_user_supabase_client)
):
    """Get document metadata by ID.

    Args:
    ----
        document_id: Document ID
        user: Current authenticated user
        supabase: Supabase client

    Returns:
    -------
        Document metadata

    """
    try:
        # Get document with case join to verify ownership
        response = (
            supabase.table("documents").select("*, cases!inner(user_id)").eq("id", document_id).execute()
        )

        if not response.data:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

        document = response.data[0]

        # Verify ownership through case
        if document["cases"]["user_id"] != user["id"]:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

        # Remove nested case data before returning
        document.pop("cases", None)

        return document
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Error fetching document: {str(e)}"
        ) from e


@router.get("/{document_id}/extracted-text")
async def get_extracted_text(
    document_id: str, user=Depends(get_current_user), supabase=Depends(get_user_supabase_client)
):
    """Get the extracted text and metadata for a specific document.

    Args:
    ----
        document_id: Document ID
        user: Current authenticated user
        supabase: Supabase client

    Returns:
    -------
        Extracted text and extraction metadata

    """
    try:
        # Get document with case join to verify ownership
        response = (
            supabase.table("documents")
            .select(
                "extracted_text, extraction_method, extraction_quality, "
                "extracted_at, page_count, ocr_provider, extraction_error, "
                "cases!inner(user_id)"
            )
            .eq("id", document_id)
            .execute()
        )

        if not response.data:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

        document = response.data[0]

        # Verify ownership through case
        if document["cases"]["user_id"] != user["id"]:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

        # Remove nested case data before returning
        document.pop("cases", None)

        return document
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error fetching extracted text: {str(e)}",
        ) from e


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    document_id: str,
    user=Depends(get_current_user),
    user_supabase=Depends(get_user_supabase_client),
    service_supabase=Depends(get_supabase_client),
):
    """Delete a document and its file from storage.

    Args:
    ----
        document_id: Document ID
        user: Current authenticated user
        user_supabase: User-scoped Supabase client
        service_supabase: Service-role Supabase client

    """
    try:
        logger.debug(f"Delete document: doc_id={document_id}, user={user['id']}")

        # Get document with ownership verification
        response = (
            user_supabase.table("documents")
            .select("storage_path, cases!inner(user_id)")
            .eq("id", document_id)
            .execute()
        )

        if not response.data:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

        document = response.data[0]

        # Verify ownership
        if document["cases"]["user_id"] != user["id"]:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

        # Delete from storage (use service client to bypass storage RLS)
        storage_path = document["storage_path"]
        logger.debug(f"Deleting from storage: {storage_path}")
        service_supabase.storage.from_("documents").remove([storage_path])

        # Delete database record (use user client for RLS)
        logger.debug("Deleting database record")
        user_supabase.table("documents").delete().eq("id", document_id).execute()

        logger.info("Document deleted successfully")
        return None
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in delete_document: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Error deleting document: {str(e)}"
        ) from e


@router.post("/bulk-delete", response_model=BulkDeleteResponse)
async def bulk_delete_documents(
    request: BulkDeleteRequest,
    user=Depends(get_current_user),
    user_supabase=Depends(get_user_supabase_client),
    service_supabase=Depends(get_supabase_client),
):
    """Bulk delete multiple documents and their files from storage.

    Args:
    ----
        request: List of document IDs to delete
        user: Current authenticated user
        user_supabase: User-scoped Supabase client
        service_supabase: Service-role Supabase client

    Returns:
    -------
        Summary of deleted and failed documents

    """
    deleted_count = 0
    failed_ids = []
    errors = []

    logger.info(f"Bulk delete request: {len(request.document_ids)} documents, user={user['id']}")

    for doc_id in request.document_ids:
        try:
            # Get document with ownership verification
            response = (
                user_supabase.table("documents")
                .select("storage_path, cases!inner(user_id)")
                .eq("id", doc_id)
                .execute()
            )

            if not response.data:
                failed_ids.append(doc_id)
                errors.append(f"Document {doc_id}: not found")
                continue

            document = response.data[0]

            # Verify ownership
            if document["cases"]["user_id"] != user["id"]:
                failed_ids.append(doc_id)
                errors.append(f"Document {doc_id}: access denied")
                continue

            # Delete from storage
            storage_path = document["storage_path"]
            try:
                service_supabase.storage.from_("documents").remove([storage_path])
            except Exception as storage_err:
                logger.warning(f"Storage deletion failed for {doc_id}: {storage_err}")
                # Continue anyway - database record should still be deleted

            # Delete database record
            user_supabase.table("documents").delete().eq("id", doc_id).execute()
            deleted_count += 1

        except Exception as e:
            logger.error(f"Error deleting document {doc_id}: {str(e)}")
            failed_ids.append(doc_id)
            errors.append(f"Document {doc_id}: {str(e)}")

    logger.info(f"Bulk delete complete: {deleted_count} deleted, {len(failed_ids)} failed")

    return BulkDeleteResponse(
        deleted_count=deleted_count,
        failed_ids=failed_ids,
        errors=errors,
    )


@router.post("/bulk-extract", response_model=BulkExtractResponse)
async def bulk_extract_documents(
    request: BulkExtractRequest,
    user=Depends(get_current_user),
    user_supabase=Depends(get_user_supabase_client),
    service_supabase=Depends(get_supabase_client),
):
    """Extract text from all documents in a case that don't have text yet."""
    import asyncio

    # Timeout per document to prevent Vercel 300s timeout
    DOC_TIMEOUT = 45  # seconds per document
    # Skip PDFs larger than this (likely to timeout on OCR)
    MAX_PDF_SIZE_FOR_BULK = 10 * 1024 * 1024  # 10MB

    try:
        logger.info(f"Bulk extraction requested for case {request.case_id} by user {user['id']}")

        # Get all documents for this case that need extraction
        # We look for documents where extracted_text is null or empty, AND status is not skipped
        response = (
            user_supabase.table("documents")
            .select("id, file_name, file_type, storage_path, file_size")
            .eq("case_id", request.case_id)
            .neq("status", DocumentStatus.SKIPPED)
            .or_("extracted_text.is.null,extracted_text.eq.''")
            .execute()
        )

        documents_to_process = response.data or []
        logger.info(f"Found {len(documents_to_process)} documents to process")

        extracted_count = 0
        failed_count = 0
        skipped_count = 0
        errors = []

        for doc in documents_to_process:
            file_name = doc.get("file_name", "unknown")
            file_type = doc.get("file_type", "")
            file_size = doc.get("file_size", 0) or 0

            # Skip large PDFs in bulk mode - they need individual processing
            is_pdf = file_type in ["application/pdf", "pdf"] or file_name.lower().endswith(".pdf")
            if is_pdf and file_size > MAX_PDF_SIZE_FOR_BULK:
                skipped_count += 1
                skip_msg = f"Skipped {file_name}: PDF too large ({file_size / (1024*1024):.1f}MB) for bulk OCR. Extract individually."
                logger.warning(skip_msg)
                errors.append(skip_msg)
                continue

            try:
                # Call trigger_extraction with timeout
                result = await asyncio.wait_for(
                    trigger_extraction(
                        document_id=doc["id"],
                        user=user,
                        user_supabase=user_supabase,
                        service_supabase=service_supabase,
                    ),
                    timeout=DOC_TIMEOUT,
                )

                if result.get("content_length", 0) > 0:
                    extracted_count += 1
                    logger.info(f"Extracted {file_name}: {result['content_length']} chars")
                else:
                    failed_count += 1
                    error_msg = f"No text extracted from {file_name} ({result.get('extraction_error', 'unsupported format')})"
                    logger.warning(error_msg)
                    errors.append(error_msg)
            except asyncio.TimeoutError:
                failed_count += 1
                error_msg = f"Timeout extracting {file_name} (>{DOC_TIMEOUT}s). Try extracting individually."
                logger.error(error_msg)
                errors.append(error_msg)
            except Exception as e:
                failed_count += 1
                error_msg = f"Failed to extract {file_name}: {str(e)}"
                logger.error(error_msg)
                errors.append(error_msg)

        return BulkExtractResponse(
            extracted_count=extracted_count,
            failed_count=failed_count + skipped_count,
            errors=errors,
        )

    except Exception as e:
        logger.error(f"Bulk extraction failed: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Bulk extraction failed: {str(e)}"
        ) from e


@router.patch("/{document_id}/verify")
async def verify_document(
    document_id: str,
    request: VerifyDocumentRequest,
    user=Depends(get_current_user),
    user_supabase=Depends(get_user_supabase_client),
):
    """Update document verification status and optionally correct extracted text.

    Args:
    ----
        document_id: Document ID
        request: Verification request with optional manual text correction
        user: Current authenticated user
        user_supabase: User-scoped Supabase client

    Returns:
    -------
        Updated document metadata

    """
    try:
        logger.debug(f"Verify document: doc_id={document_id}, user={user['id']}")

        # Get document with ownership verification
        response = (
            user_supabase.table("documents")
            .select("id, extracted_text, manual_text, cases!inner(user_id)")
            .eq("id", document_id)
            .execute()
        )

        if not response.data:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

        document = response.data[0]

        # Verify ownership
        if document["cases"]["user_id"] != user["id"]:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

        # Validate that document has text if being verified as ready
        if request.is_verified and not (document.get("extracted_text") or document.get("manual_text") or request.manual_text):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Cannot verify document without extracted text. Please run OCR first.",
            )

        # Build update payload
        update_data = {
            "is_verified": request.is_verified,
            "is_flagged_as_junk": request.is_flagged_as_junk,
            "status": DocumentStatus.READY if request.is_verified else DocumentStatus.NEEDS_REVIEW,
            "updated_at": datetime.utcnow().isoformat(),
        }

        if request.manual_text is not None:
            update_data["manual_text"] = request.manual_text
            update_data["text_edited_at"] = datetime.utcnow().isoformat()

        # Update document
        update_response = user_supabase.table("documents").update(update_data).eq("id", document_id).execute()

        if not update_response.data:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to update document",
            )

        logger.info(f"Document {document_id} verified successfully")
        return update_response.data[0]

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in verify_document: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error verifying document: {str(e)}",
        ) from e


class ToggleExclusionRequest(BaseModel):
    """Request model for toggling document exclusion."""

    excluded: bool


@router.patch("/{document_id}/exclusion")
async def toggle_document_exclusion(
    document_id: str,
    request: ToggleExclusionRequest,
    user=Depends(get_current_user),
    user_supabase=Depends(get_user_supabase_client),
):
    """Toggle document exclusion status (for duplicate management).

    Allows users to include/exclude duplicate documents from analysis.

    Args:
    ----
        document_id: Document ID
        request: Exclusion status to set
        user: Current authenticated user
        user_supabase: User-scoped Supabase client

    Returns:
    -------
        Updated document with new exclusion status

    """
    try:
        logger.debug(f"Toggle exclusion: doc_id={document_id}, excluded={request.excluded}, user={user['id']}")

        # Get document with ownership verification
        response = (
            user_supabase.table("documents")
            .select("id, metadata, cases!inner(user_id)")
            .eq("id", document_id)
            .execute()
        )

        if not response.data:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

        document = response.data[0]

        # Verify ownership
        if document["cases"]["user_id"] != user["id"]:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

        # Update metadata with new exclusion status
        metadata = document.get("metadata", {}) or {}
        metadata["excluded"] = request.excluded

        # Update status based on exclusion
        new_status = "duplicate" if metadata.get("is_duplicate") and request.excluded else "ready"

        update_data = {
            "metadata": metadata,
            "status": new_status,
            "updated_at": datetime.utcnow().isoformat(),
        }

        user_supabase.table("documents").update(update_data).eq("id", document_id).execute()

        logger.info(f"Document exclusion toggled: {document_id}, excluded={request.excluded}")

        return {
            "document_id": document_id,
            "excluded": request.excluded,
            "status": new_status,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error toggling document exclusion: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error toggling exclusion: {str(e)}",
        ) from e


def is_photo_rejection_message(text: str) -> bool:
    """Check if OCR result is a message indicating the image is a photo, not a text document."""
    if not text or len(text) < 20:
        return False
    
    rejection_phrases = [
        "unable to extract text",
        "unable to access",
        "unable to analyze",
        "can't access",
        "can't analyze",
        "cannot access",
        "cannot analyze",
        "not a legal document",
        "appears to be a photo",
        "this is a photo",
        "this image is",
        "does not contain",
        "no text to extract",
        "cannot extract text",
        "I can't extract",
        "I'm unable to extract",
        "I'm unable to access",
        "I'm unable to analyze",
        "how to describe an image",
        "help you understand how to describe",
        "general guide on how",
    ]
    
    text_lower = text.lower()
    return any(phrase in text_lower for phrase in rejection_phrases)


async def get_case_context(case_id: str, supabase_client) -> dict:
    """Get case context for image analysis."""
    try:
        response = supabase_client.table("cases").select("case_name, client_name, description").eq("id", case_id).execute()
        if response.data and len(response.data) > 0:
            case = response.data[0]
            return {
                "case_name": case.get("case_name", "Unknown Case"),
                "client_name": case.get("client_name", "Unknown Client"),
                "description": case.get("description", ""),
            }
    except Exception as e:
        logger.warning(f"Could not fetch case context for {case_id}: {e}")
    
    return {"case_name": "Legal Case", "client_name": "", "description": ""}


async def analyze_image_with_vision(file_bytes: bytes, file_name: str, case_context: dict) -> tuple[str, str]:
    """
    Analyze an image using vision AI with case context.
    
    Returns (visual_description, extraction_method)
    """
    import asyncio
    import base64
    from starlette.concurrency import run_in_threadpool
    from legal_portal.utils.openai_client import OpenAIClient
    
    try:
        logger.info(f"Analyzing image content for {file_name} with case context")
        openai_client = OpenAIClient()
        client = openai_client.client
        
        # Determine MIME type
        mime_type = "image/jpeg" if file_name.lower().endswith((".jpg", ".jpeg")) else "image/png"
        base64_image = base64.b64encode(file_bytes).decode("utf-8")
        
        # Build context-aware prompt
        case_info = f"Case: {case_context['case_name']}"
        if case_context.get('client_name'):
            case_info += f" (Client: {case_context['client_name']})"
        if case_context.get('description'):
            case_info += f"\nCase Description: {case_context['description']}"
        
        prompt = (
            f"You are analyzing an image as photographic evidence for a legal case.\n\n"
            f"CASE CONTEXT:\n{case_info}\n\n"
            f"IMAGE FILE: {file_name}\n\n"
            "TASK: Describe exactly what you see in this image. This is actual photographic evidence, "
            "not a document to extract text from.\n\n"
            "In your description, include:\n"
            "1. OBJECTS & SCENE: What physical objects, people, locations, or structures are visible\n"
            "2. CONDITIONS: Any damage, defects, injuries, wear, contamination, or unusual conditions\n"
            "3. VISIBLE TEXT: Any text, labels, signs, dates, numbers, or markings you can see\n"
            "4. SETTING: The environment/location context (indoor/outdoor, type of space, lighting)\n"
            "5. CASE RELEVANCE: How this visual evidence relates to the legal matter described above\n\n"
            "IMPORTANT: DO NOT provide instructions or templates. Describe the ACTUAL SPECIFIC image "
            "you are viewing right now. Start your response by describing what you see."
        )
        
        def vision_analysis():
            return client.chat.completions.create(
                model="gpt-4o",
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{base64_image}",
                                "detail": "high"
                            }
                        },
                    ]
                }],
                max_tokens=1500,
                temperature=0.3,
            )
        
        response = await asyncio.wait_for(
            run_in_threadpool(vision_analysis),
            timeout=60.0,
        )
        
        description = response.choices[0].message.content
        logger.info(f"Successfully analyzed image {file_name} with vision AI: {len(description)} chars")
        return description, "GPT-4o Vision (Image Analysis)"
        
    except Exception as e:
        logger.error(f"Vision analysis failed for {file_name}: {e}")
        raise


@router.post("/{document_id}/extract")
async def trigger_extraction(
    document_id: str,
    user=Depends(get_current_user),
    user_supabase=Depends(get_user_supabase_client),
    service_supabase=Depends(get_supabase_client),
):
    """Manually trigger or re-run text extraction for a single document.

    Args:
    ----
        document_id: Document ID
        user: Current authenticated user
        user_supabase: User-scoped Supabase client
        service_supabase: Service-role Supabase client

    Returns:
    -------
        Extraction result with extracted text and metadata

    """
    import os
    import tempfile

    from legal_portal.core.data_models import DocumentType
    from legal_portal.services.file_processors.pdf_processor import process_pdf

    try:
        logger.info(f"Trigger extraction: doc_id={document_id}, user={user['id']}")

        # Get document with ownership verification
        response = (
            user_supabase.table("documents").select("*, cases!inner(user_id)").eq("id", document_id).execute()
        )

        if not response.data:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

        document = response.data[0]

        # Verify ownership
        if document["cases"]["user_id"] != user["id"]:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

        # Download file from storage
        storage_path = document["storage_path"]
        logger.debug(f"Downloading file from storage: {storage_path}")

        file_bytes = service_supabase.storage.from_("documents").download(storage_path)

        if not file_bytes:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to download document from storage",
            )

        # Determine file type and process
        file_name = document["file_name"]
        file_type = document["file_type"]

        extracted_text = ""
        extraction_method = ""
        extraction_quality = "high"
        ocr_provider = None
        extraction_error = None
        page_count = None

        if file_type in ["application/pdf", "pdf"]:
            # Write to temp file for PDF processing
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name

            try:
                result = await process_pdf(
                    file_path=tmp_path,
                    document_type=DocumentType.CASE_DOCUMENT,
                    original_filename=file_name,
                )
                extracted_text = result.content
                extraction_method = result.extraction_method or "unknown"
                extraction_quality = result.extraction_quality or "high"
                ocr_provider = result.ocr_provider
                extraction_error = result.extraction_error
                page_count = result.page_count
            finally:
                os.unlink(tmp_path)

        elif file_type in ["text/plain", "txt"]:
            # Plain text file
            extracted_text = file_bytes.decode("utf-8", errors="replace")
            extraction_method = "direct_text"
            extraction_quality = "high"

        elif file_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
            "docx",
            "doc",
        ] or file_name.lower().endswith((".docx", ".doc")):
            # Microsoft Word document - extract text directly (no OCR needed)
            import io
            try:
                import docx
                document = docx.Document(io.BytesIO(file_bytes))

                # Extract text from paragraphs
                paragraphs = [para.text for para in document.paragraphs if para.text.strip()]

                # Also extract text from tables
                table_text = []
                for table in document.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_text.append(" | ".join(row_text))

                # Combine paragraphs and table content
                all_text = paragraphs + table_text
                extracted_text = "\n".join(all_text)

                extraction_method = "docx_direct"
                extraction_quality = "high" if len(extracted_text) > 50 else "medium"
                logger.info(f"DOCX extraction: {len(paragraphs)} paragraphs, {len(table_text)} table rows")

            except ImportError:
                extraction_error = "python-docx library not available for DOCX extraction"
                extraction_method = "none"
                extraction_quality = "low"
            except Exception as docx_err:
                extraction_error = f"DOCX extraction failed: {str(docx_err)}"
                extraction_method = "none"
                extraction_quality = "low"
                logger.error(f"DOCX extraction error for {file_name}: {docx_err}")

        elif file_type in ["image/png", "image/jpeg", "image/jpg"] or file_name.lower().endswith((".png", ".jpg", ".jpeg")):
            # Image file - use Google Vision OCR (with GPT-4o fallback)
            try:
                google_client = GoogleVisionClient.get_instance()
                if google_client.is_available:
                    try:
                        import asyncio

                        from starlette.concurrency import run_in_threadpool

                        def do_google_ocr():
                            return google_client.extract_text_from_image(file_bytes)

                        vision_text = await asyncio.wait_for(
                            run_in_threadpool(do_google_ocr),
                            timeout=30.0,
                        )

                        if vision_text and vision_text.strip():
                            extracted_text = vision_text
                            extraction_method = "Google Cloud Vision"
                            extraction_quality = "high"
                            ocr_provider = "google_vision"
                            logger.info(f"Successfully extracted text from {file_name} using Google Vision")
                            
                            # Check if the OCR result is actually a photo rejection message
                            if is_photo_rejection_message(extracted_text):
                                logger.info(f"Google Vision detected photo (non-text image) in {file_name}, switching to visual analysis")
                                try:
                                    # Get case context
                                    case_context = await get_case_context(document["case_id"], user_supabase)
                                    
                                    # Analyze image with context
                                    visual_description, vision_method = await analyze_image_with_vision(
                                        file_bytes, file_name, case_context
                                    )
                                    
                                    if visual_description and len(visual_description) > 50:
                                        extracted_text = visual_description
                                        extraction_method = vision_method
                                        extraction_quality = "high"
                                        ocr_provider = "openai_vision_analysis"
                                        
                                        # Mark as visual content in metadata
                                        document_metadata = document.get("metadata", {}) or {}
                                        document_metadata["is_visual_content"] = True
                                        
                                        # Update metadata in database
                                        user_supabase.table("documents").update({
                                            "metadata": document_metadata
                                        }).eq("id", document_id).execute()
                                        
                                        logger.info(f"Successfully analyzed photo content for {file_name}")
                                    else:
                                        logger.warning(f"Vision analysis returned insufficient content for {file_name}")
                                except Exception as vision_err:
                                    logger.error(f"Vision analysis failed for {file_name}: {vision_err}")
                                    # Keep the original rejection message if vision analysis fails
                        else:
                            raise ValueError("Google Vision returned empty text")
                    except Exception as google_err:
                        logger.warning(f"{file_name}: Google Vision failed ({google_err}). Falling back to GPT-4o Vision.")
                        raise  # Fall through to GPT-4o fallback
                else:
                    raise ValueError("Google Vision client not available")
            except Exception:
                # Fallback to GPT-4o Vision
                try:
                    import asyncio
                    import base64

                    from starlette.concurrency import run_in_threadpool

                    from legal_portal.utils.openai_client import OpenAIClient

                    logger.info(f"Using GPT-4o Vision for {file_name}")
                    openai_client = OpenAIClient()
                    client = openai_client.client

                    # Determine MIME type
                    if file_type in ["image/jpeg", "image/jpg"] or file_name.lower().endswith((".jpg", ".jpeg")):
                        mime_type = "image/jpeg"
                    else:
                        mime_type = "image/png"

                    base64_image = base64.b64encode(file_bytes).decode("utf-8")
                    prompt = (
                        f"Extract ALL text from this legal document image. "
                        f"Filename: {file_name}. "
                        "This is a scanned document that needs OCR text extraction. "
                        "Maintain the logical structure and layout. "
                        "If there are tables, preserve the row/column relationship. "
                        "Provide the text verbatim including all numbers, dates, and names. Do not summarize."
                    )

                    def gpt4o_ocr():
                        return client.chat.completions.create(
                            model="gpt-4o",
                            messages=[{
                                "role": "user",
                                "content": [
                                    {"type": "text", "text": prompt},
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:{mime_type};base64,{base64_image}",
                                            "detail": "high"
                                        }
                                    },
                                ]
                            }],
                            max_tokens=4000,
                            temperature=0.0,
                        )

                    response = await asyncio.wait_for(
                        run_in_threadpool(gpt4o_ocr),
                        timeout=60.0,
                    )

                    extracted_text = response.choices[0].message.content
                    extraction_method = "GPT-4o Vision"
                    extraction_quality = "high"
                    ocr_provider = "openai"
                    logger.info(f"Successfully extracted text from {file_name} using GPT-4o Vision")
                    
                    # Check if the OCR result is actually a photo rejection message
                    if is_photo_rejection_message(extracted_text):
                        logger.info(f"Detected photo (non-text image) in {file_name}, switching to visual analysis")
                        try:
                            # Get case context
                            case_context = await get_case_context(document["case_id"], user_supabase)
                            
                            # Analyze image with context
                            visual_description, vision_method = await analyze_image_with_vision(
                                file_bytes, file_name, case_context
                            )
                            
                            if visual_description and len(visual_description) > 50:
                                extracted_text = visual_description
                                extraction_method = vision_method
                                extraction_quality = "high"
                                ocr_provider = "openai_vision_analysis"
                                
                                # Mark as visual content in metadata
                                document_metadata = document.get("metadata", {}) or {}
                                document_metadata["is_visual_content"] = True
                                
                                # Update metadata in database
                                user_supabase.table("documents").update({
                                    "metadata": document_metadata
                                }).eq("id", document_id).execute()
                                
                                logger.info(f"Successfully analyzed photo content for {file_name}")
                            else:
                                logger.warning(f"Vision analysis returned insufficient content for {file_name}")
                        except Exception as vision_err:
                            logger.error(f"Vision analysis failed for {file_name}: {vision_err}")
                            # Keep the original rejection message if vision analysis fails
                    
                except Exception as ocr_err:
                    extraction_error = f"Image OCR failed: {str(ocr_err)}"
                    extraction_method = "failed"
                    extraction_quality = "low"
                    logger.error(f"Image extraction error for {file_name}: {ocr_err}")

        elif file_type in ["message/rfc822", "eml"] or file_name.lower().endswith(".eml"):
            # Email file (.eml)
            try:
                import tempfile

                # Write to temp file for eml processing
                with tempfile.NamedTemporaryFile(suffix=".eml", delete=False) as tmp:
                    tmp.write(file_bytes)
                    tmp_path = tmp.name

                try:
                    # Process eml file asynchronously
                    result = await process_eml(
                        file_path=tmp_path,
                        document_type=DocumentType.CORRESPONDENCE,
                        original_filename=file_name,
                    )

                    extracted_text = result.content
                    extraction_method = result.extraction_method or "email_parser"
                    extraction_quality = result.extraction_quality or "high"
                    extraction_error = result.extraction_error

                    logger.info(f"Successfully extracted email content from {file_name}: {len(extracted_text)} chars")
                finally:
                    # Clean up temp file
                    import os
                    try:
                        os.unlink(tmp_path)
                    except Exception:
                        pass
            except Exception as e:
                extraction_error = f"Email extraction failed: {str(e)}"
                extraction_method = "failed"
                extraction_quality = "low"
                logger.error(f"Email extraction error for {file_name}: {e}")

        else:
            # Unsupported type
            extraction_error = f"Unsupported file type for extraction: {file_type}"
            extraction_method = "none"
            extraction_quality = "low"

        # Sanitize extracted text to remove NULL characters that PostgreSQL can't store
        extracted_text = sanitize_text_for_db(extracted_text)

        # Update document with extraction results
        update_data = {
            "extracted_text": extracted_text,
            "extraction_method": extraction_method,
            "extraction_quality": extraction_quality,
            "ocr_provider": ocr_provider,
            "extraction_error": extraction_error,
            "page_count": page_count,
            "extracted_at": datetime.utcnow().isoformat() if extracted_text else None,
            "updated_at": datetime.utcnow().isoformat(),
            "status": DocumentStatus.READY if extracted_text and extracted_text.strip() else DocumentStatus.EXTRACTION_FAILED,
        }

        update_result = user_supabase.table("documents").update(update_data).eq("id", document_id).execute()

        # CRITICAL: Verify the update actually saved to the database
        if not update_result.data:
            logger.error(f"Failed to save extracted text for document {document_id} - update returned no data")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to save extraction results to database. Please try again.",
            )

        logger.info(
            f"Extraction complete for {document_id}: method={extraction_method}, "
            f"quality={extraction_quality}, chars={len(extracted_text or '')}"
        )

        return {
            "document_id": document_id,
            "extracted_text": extracted_text,
            "extraction_method": extraction_method,
            "extraction_quality": extraction_quality,
            "ocr_provider": ocr_provider,
            "extraction_error": extraction_error,
            "page_count": page_count,
            "content_length": len(extracted_text),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in trigger_extraction: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error extracting text: {str(e)}",
        ) from e


@router.post("/{document_id}/replace", response_model=DocumentResponse)
async def replace_document_file(
    document_id: str,
    file: UploadFile = File(...),
    user=Depends(get_current_user),
    user_supabase=Depends(get_user_supabase_client),
    service_supabase=Depends(get_supabase_client),
):
    """Replace the file for an existing document record.

    Useful for fixing documents with download errors or corruption.
    """
    import os
    import tempfile

    from legal_portal.services.file_processors.pdf_processor import process_pdf

    try:
        logger.info(f"Replacing file for document {document_id}")

        # Get existing document with ownership verification
        response = (
            user_supabase.table("documents")
            .select("id, case_id, cases!inner(user_id), storage_path")
            .eq("id", document_id)
            .execute()
        )

        if not response.data:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

        existing_doc = response.data[0]
        case_id = existing_doc["case_id"]

        # Verify ownership
        if existing_doc["cases"]["user_id"] != user["id"]:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

        # Read new file content
        file_content = await file.read()

        # Delete old file from storage if it exists
        if existing_doc.get("storage_path"):
            try:
                service_supabase.storage.from_("documents").remove([existing_doc["storage_path"]])
            except Exception as e:
                logger.warning(f"Failed to delete old file {existing_doc['storage_path']}: {e}")

        # Use unified processor for validation and upload
        processor = DocumentProcessor()
        doc_record = await processor.process_and_upload(
            file_content=file_content,
            filename=file.filename,
            user_id=user["id"],
            case_id=case_id,
            supabase_client=service_supabase,
            is_intake_form=False,  # Can be adjusted if needed
            content_type=file.content_type,
        )

        # Extract text from the new file
        extracted_text = ""
        extraction_method = ""
        extraction_quality = "high"
        ocr_provider = None
        extraction_error = None
        page_count = None

        file_type = doc_record.get("file_type", file.content_type)
        file_name = doc_record.get("file_name", file.filename)

        if file_type in ["application/pdf", "pdf"] or file_name.lower().endswith(".pdf"):
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name

            try:
                result = await process_pdf(
                    file_path=tmp_path,
                    document_type=DocumentType.CASE_DOCUMENT,
                    original_filename=file_name,
                )
                extracted_text = result.content
                extraction_method = result.extraction_method or "unknown"
                extraction_quality = result.extraction_quality or "high"
                ocr_provider = result.ocr_provider
                extraction_error = result.extraction_error
                page_count = result.page_count
            finally:
                os.unlink(tmp_path)
        elif file_type in ["text/plain", "txt"] or file_name.lower().endswith(".txt"):
            try:
                extracted_text = file_content.decode("utf-8", errors="replace")
                extraction_method = "direct_text"
                extraction_quality = "high"
            except Exception as e:
                extraction_error = f"Failed to decode text: {e}"
                extraction_method = "failed"
                extraction_quality = "low"

        elif file_type in ["message/rfc822", "eml"] or file_name.lower().endswith(".eml"):
            # Email file (.eml)
            try:
                import tempfile

                # Write to temp file for eml processing
                with tempfile.NamedTemporaryFile(suffix=".eml", delete=False) as tmp:
                    tmp.write(file_content)
                    tmp_path = tmp.name

                try:
                    # Process eml file asynchronously
                    result = await process_eml(
                        file_path=tmp_path,
                        document_type=DocumentType.CORRESPONDENCE,
                        original_filename=file_name,
                    )

                    extracted_text = result.content
                    extraction_method = result.extraction_method or "email_parser"
                    extraction_quality = result.extraction_quality or "high"
                    extraction_error = result.extraction_error

                    logger.info(f"Successfully extracted email content from {file_name}: {len(extracted_text)} chars")
                finally:
                    # Clean up temp file
                    import os
                    try:
                        os.unlink(tmp_path)
                    except Exception:
                        pass
            except Exception as e:
                extraction_error = f"Email extraction failed: {str(e)}"
                extraction_method = "failed"
                extraction_quality = "low"
                logger.error(f"Email extraction error for {file_name}: {e}")

        # Sanitize extracted text to remove NULL characters that PostgreSQL can't store
        extracted_text = sanitize_text_for_db(extracted_text)

        # Update existing document record
        update_data = {
            "file_name": file_name,
            "file_type": file_type,
            "file_size": doc_record["file_size"],
            "storage_path": doc_record["storage_path"],
            "status": DocumentStatus.READY if extracted_text else DocumentStatus.EXTRACTION_FAILED,
            "extracted_text": extracted_text if extracted_text else None,
            "extraction_method": extraction_method,
            "extraction_quality": extraction_quality,
            "ocr_provider": ocr_provider,
            "extraction_error": extraction_error,
            "page_count": page_count,
            "metadata": {**existing_doc.get("metadata", {}), **doc_record["metadata"]},
            "updated_at": datetime.utcnow().isoformat(),
        }

        update_response = user_supabase.table("documents").update(update_data).eq("id", document_id).execute()

        if not update_response.data:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to update document record"
            )

        return update_response.data[0]

    except HTTPException:
        raise
    except ValidationError as e:
        return JSONResponse(
            status_code=status.HTTP_400_BAD_REQUEST,
            content={"code": e.error_code, "detail": str(e)},
        )
    except Exception as e:
        logger.error(f"Error in replace_document_file: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Error replacing document: {str(e)}"
        ) from e
