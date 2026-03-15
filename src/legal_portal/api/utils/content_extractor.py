"""Document processing utilities for downloading and extracting text from files."""

import io
import re
import time
import zipfile
import xml.etree.ElementTree as ET
from typing import Optional, Tuple

import requests

from legal_portal.services.file_compression_service import get_compression_service
from legal_portal.utils.compression_utils import format_file_size
from legal_portal.utils.logging_config import get_module_logger

logger = get_module_logger(__name__)

# Rate limiting for Clio API: 30 requests per 10 seconds = ~3 requests/second
# Use 0.4s delay to stay safely under the limit
CLIO_RATE_LIMIT_DELAY = 0.4


class ClioRateLimiter:
    """Encapsulates Clio API rate-limiting state."""

    def __init__(self, delay: float = CLIO_RATE_LIMIT_DELAY):
        self._last_request_time = 0.0
        self._delay = delay

    def wait_if_needed(self):
        """Sleep if needed to respect rate limit, then record the request time."""
        elapsed = time.time() - self._last_request_time
        if elapsed < self._delay:
            sleep_time = self._delay - elapsed
            logger.debug(f"Rate limiting: sleeping {sleep_time:.2f}s before Clio request")
            time.sleep(sleep_time)
        self._last_request_time = time.time()


_clio_rate_limiter = ClioRateLimiter()

# Conditional imports for PDF extraction
# Try pypdf first (lightweight, works on Vercel), then fitz (PyMuPDF, better quality)
PYPDF_AVAILABLE = False
FITZ_AVAILABLE = False

try:
    from pypdf import PdfReader

    PYPDF_AVAILABLE = True
    logger.debug("pypdf available for PDF extraction")
except ImportError:
    logger.debug("pypdf not available")

try:
    import fitz  # PyMuPDF

    FITZ_AVAILABLE = True
    logger.debug("PyMuPDF (fitz) available for PDF extraction")
except ImportError:
    logger.debug("PyMuPDF (fitz) not available")

# Conditional import for DOCX
DOCX_AVAILABLE = False
try:
    from docx import Document

    DOCX_AVAILABLE = True
    logger.debug("python-docx available for DOCX extraction")
except ImportError:
    logger.debug("python-docx not available")


class DocumentProcessor:
    """Handles document download and text extraction."""

    _WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    @staticmethod
    def download_file(url: str, access_token: str, max_retries: int = 5) -> Tuple[bytes, str]:
        """Download a file from a URL with authentication.
        
        Includes rate limiting and retry logic with exponential backoff for 429 errors.

        Args:
        ----
            url: URL to download from
            access_token: OAuth access token for authentication
            max_retries: Maximum number of retry attempts (default: 5)

        Returns:
        -------
            Tuple of (file_content, content_type)

        Raises:
        ------
            Exception: If download fails after all retries

        """
        headers = {
            "Authorization": f"Bearer {access_token}",
        }

        # Check if this is a Clio API request
        is_clio_request = "app.clio.com" in url

        for attempt in range(max_retries):
            try:
                # Rate limiting for Clio API
                if is_clio_request:
                    _clio_rate_limiter.wait_if_needed()

                response = requests.get(url, headers=headers, timeout=60)

                # Handle 429 Too Many Requests with exponential backoff
                if response.status_code == 429:
                    if attempt < max_retries - 1:
                        # Exponential backoff: 2^attempt seconds (1s, 2s, 4s, 8s, 16s)
                        wait_time = 2 ** attempt
                        # Also check Retry-After header if present
                        retry_after = response.headers.get("Retry-After")
                        if retry_after:
                            try:
                                wait_time = max(wait_time, int(retry_after))
                            except ValueError:
                                pass
                        logger.warning(
                            f"Rate limited (429) on attempt {attempt + 1}/{max_retries}. "
                            f"Waiting {wait_time}s before retry..."
                        )
                        time.sleep(wait_time)
                        continue
                    else:
                        raise Exception(f"Rate limit exceeded after {max_retries} attempts")

                response.raise_for_status()

                content_type = response.headers.get("content-type", "application/octet-stream")
                return response.content, content_type

            except requests.exceptions.RequestException as e:
                if attempt < max_retries - 1:
                    # For other errors, use shorter backoff
                    wait_time = min(2 ** attempt, 5)  # Cap at 5 seconds
                    logger.warning(
                        f"Request failed on attempt {attempt + 1}/{max_retries}: {e}. "
                        f"Retrying in {wait_time}s..."
                    )
                    time.sleep(wait_time)
                    continue
                else:
                    raise Exception(f"Download failed after {max_retries} attempts: {e}") from e

        raise Exception(f"Download failed after {max_retries} attempts")

    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file.

        Uses pypdf (lightweight) or falls back to PyMuPDF (fitz) if available.

        Args:
        ----
            file_content: PDF file bytes

        Returns:
        -------
            Extracted text

        """
        # Try PyMuPDF first (better quality extraction) if available
        if FITZ_AVAILABLE:
            try:
                pdf_document = fitz.open(stream=file_content, filetype="pdf")
                text_parts = []
                for page_num in range(pdf_document.page_count):
                    page = pdf_document[page_num]
                    text_parts.append(page.get_text())
                pdf_document.close()
                return "\n\n".join(text_parts)
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed, trying pypdf: {e}")

        # Fall back to pypdf (lightweight, works on Vercel)
        if PYPDF_AVAILABLE:
            try:
                reader = PdfReader(io.BytesIO(file_content))
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return "\n\n".join(text_parts)
            except Exception as e:
                raise Exception(f"Failed to extract text from PDF with pypdf: {str(e)}") from e

        # No PDF library available
        raise Exception("No PDF extraction library available (install pypdf or PyMuPDF)")

    @staticmethod
    def _extract_word_paragraph_text(paragraph: ET.Element) -> str:
        """Extract readable text from a WordprocessingML paragraph element."""
        ns = f"{{{DocumentProcessor._WORD_NS}}}"
        chunks = []
        for element in paragraph.iter():
            if element.tag == f"{ns}t" and element.text:
                chunks.append(element.text)
            elif element.tag == f"{ns}tab":
                chunks.append("\t")
            elif element.tag in (f"{ns}br", f"{ns}cr"):
                chunks.append("\n")
        return "".join(chunks).strip()

    @staticmethod
    def _extract_text_from_docx_xml(file_content: bytes) -> str:
        """Extract DOCX text directly from zipped WordprocessingML XML."""
        ns = f"{{{DocumentProcessor._WORD_NS}}}"
        try:
            with zipfile.ZipFile(io.BytesIO(file_content)) as archive:
                if "word/document.xml" not in archive.namelist():
                    raise Exception("Missing word/document.xml")

                lines: list[str] = []
                root = ET.fromstring(archive.read("word/document.xml"))
                body = root.find(f".//{ns}body")
                if body is None:
                    raise Exception("Malformed DOCX: missing document body")

                for child in body:
                    if child.tag == f"{ns}p":
                        paragraph_text = DocumentProcessor._extract_word_paragraph_text(child)
                        if paragraph_text:
                            lines.append(paragraph_text)
                    elif child.tag == f"{ns}tbl":
                        for row in child.findall(f"{ns}tr"):
                            row_cells = []
                            for cell in row.findall(f"{ns}tc"):
                                cell_paragraphs = []
                                for paragraph in cell.findall(f".//{ns}p"):
                                    paragraph_text = DocumentProcessor._extract_word_paragraph_text(paragraph)
                                    if paragraph_text:
                                        cell_paragraphs.append(paragraph_text)
                                cell_text = " ".join(cell_paragraphs).strip()
                                if cell_text:
                                    row_cells.append(cell_text)
                            if row_cells:
                                lines.append(" | ".join(row_cells))

                return "\n".join(lines)
        except zipfile.BadZipFile as e:
            raise Exception("Invalid DOCX container (not a valid ZIP)") from e
        except ET.ParseError as e:
            raise Exception("Invalid DOCX XML structure") from e

    @staticmethod
    def extract_text_from_docx_with_method(file_content: bytes) -> Tuple[str, str]:
        """Extract text from DOCX and return both content and extraction backend."""
        if DOCX_AVAILABLE:
            try:
                doc = Document(io.BytesIO(file_content))
                paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
                table_rows = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_rows.append(" | ".join(row_text))

                text = "\n".join(paragraphs + table_rows).strip()
                if text:
                    return text, "python-docx"
                logger.warning("python-docx returned empty content; trying XML fallback")
            except Exception as e:
                logger.warning(f"python-docx extraction failed, trying XML fallback: {e}")

        try:
            text = DocumentProcessor._extract_text_from_docx_xml(file_content).strip()
            if text:
                return text, "docx_xml_fallback"
            raise Exception("DOCX XML extraction returned empty content")
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}") from e

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file.

        Args:
        ----
            file_content: DOCX file bytes

        Returns:
        -------
            Extracted text

        """
        extracted_text, _ = DocumentProcessor.extract_text_from_docx_with_method(file_content)
        return extracted_text

    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        """Extract text from plain text file.

        Args:
        ----
            file_content: Text file bytes

        Returns:
        -------
            Extracted text

        """
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                return file_content.decode("utf-8")
            except UnicodeDecodeError:
                return file_content.decode("latin-1", errors="replace")
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {str(e)}") from e

    @classmethod
    def extract_text(cls, file_content: bytes, content_type: str, filename: str = "") -> Optional[str]:
        """Extract text from file based on content type.

        Args:
        ----
            file_content: File bytes
            content_type: MIME type of the file
            filename: Original filename (used for extension fallback)

        Returns:
        -------
            Extracted text or None if extraction not supported

        """
        # Normalize content type
        content_type = content_type.lower().split(";")[0].strip()

        # Check file extension as fallback
        extension = ""
        if filename and "." in filename:
            extension = filename.split(".")[-1].lower()

        extracted = None

        # PDF
        if content_type == "application/pdf" or extension == "pdf":
            if not PYPDF_AVAILABLE and not FITZ_AVAILABLE:
                logger.warning(f"No PDF library available to extract text from {filename}")
                return None
            extracted = cls.extract_text_from_pdf(file_content)

        # DOCX
        elif (
            content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or extension == "docx"
        ):
            extracted = cls.extract_text_from_docx(file_content)

        # Plain text
        elif content_type.startswith("text/") or extension in ["txt", "text", "log", "md"]:
            extracted = cls.extract_text_from_txt(file_content)
        # Unsupported type
        else:
            return None

        # Sanitize extracted text to remove NULL characters that PostgreSQL can't store
        if extracted:
            # Remove NULL characters (\\x00 and \\u0000)
            extracted = extracted.replace('\x00', '')
            # Remove other problematic control characters (except newline, tab, carriage return)
            extracted = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f]', '', extracted)

        return extracted

    @classmethod
    def download_and_extract(
        cls, url: str, access_token: str, filename: str = "", compress: bool = True
    ) -> Tuple[bytes, str, Optional[str], Optional[dict]]:
        """Download file, optionally compress it, and extract text in one operation.

        Args:
        ----
            url: URL to download from
            access_token: OAuth access token
            filename: Original filename for type detection
            compress: Whether to attempt compression for large files (default: True)

        Returns:
        -------
            Tuple of (file_content, content_type, extracted_text, compression_metadata)
            compression_metadata includes: {
                "compressed": bool,
                "original_size": int,
                "compressed_size": int,
                "compression_ratio": float,
                "method": str
            }

        Raises:
        ------
            Exception: If download fails

        """
        # Download file
        file_content, content_type = cls.download_file(url, access_token)
        original_size = len(file_content)

        # Log download
        logger.info(f"Downloaded file: {filename} ({format_file_size(original_size)})")

        # Initialize compression metadata
        compression_metadata = {
            "compressed": False,
            "original_size": original_size,
            "compressed_size": original_size,
            "compression_ratio": 1.0,
            "method": "none",
        }

        # Attempt compression if enabled and file is large enough
        if compress:
            try:
                compression_service = get_compression_service()
                compression_result = compression_service.compress_file(file_content, filename, content_type)

                # Update file content if compressed
                if compression_result.was_compressed:
                    file_content = compression_result.compressed_data
                    logger.info(
                        f"Compression applied: {format_file_size(original_size)} → "
                        f"{format_file_size(compression_result.compressed_size)} "
                        f"({compression_result.method_used})"
                    )

                # Update compression metadata
                compression_metadata = {
                    "compressed": compression_result.was_compressed,
                    "original_size": compression_result.original_size,
                    "compressed_size": compression_result.compressed_size,
                    "compression_ratio": compression_result.compression_ratio,
                    "method": compression_result.method_used,
                }

            except Exception as e:
                logger.warning(f"Compression attempt failed for {filename}: {e}")
                # Continue with uncompressed file

        # Extract text (may return None for unsupported types)
        # Note: Extract from the potentially compressed file
        try:
            extracted_text = cls.extract_text(file_content, content_type, filename)
        except Exception as e:
            logger.warning(f"Text extraction failed for {filename}: {e}")
            extracted_text = None

        return file_content, content_type, extracted_text, compression_metadata
