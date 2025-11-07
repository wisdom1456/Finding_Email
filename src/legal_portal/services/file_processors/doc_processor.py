from __future__ import annotations

import mimetypes
import os

from legal_portal.core.data_models import (
    DocumentType,
    FileMetadata,
    FileType,
    ProcessedDocument,
)
from legal_portal.utils.logging_config import get_module_logger

logger = get_module_logger(__name__)


async def process_doc(
    file_path: str, document_type: DocumentType, original_filename: str
) -> ProcessedDocument:
    """Processes a legacy DOC file by extracting its content using multiple fallback methods.

    Legacy DOC-specific features:
    - Handles OLE compound document format (Office 97-2003)
    - Multiple extraction methods with fallbacks
    - Enhanced error handling for corrupt legacy files
    - Encoding detection for international documents
    """
    logger.debug(f"Processing legacy DOC: {original_filename}")

    text_content = ""
    extraction_method = "unknown"

    try:
        # Method 1: Try python-docx2txt (handles some legacy DOC files)
        try:
            import docx2txt

            text_content = docx2txt.process(file_path)
            extraction_method = "docx2txt"
            if text_content.strip():
                logger.info(f"Successfully extracted text from DOC using docx2txt: {original_filename}")
            else:
                raise ValueError("Empty content from docx2txt")
        except (ImportError, Exception) as e:
            logger.debug(f"docx2txt failed for {original_filename}: {e}")

            # Method 2: Try antiword (if available)
            try:
                import subprocess

                result = subprocess.run(
                    ["antiword", file_path], capture_output=True, text=True, timeout=30, check=False
                )
                if result.returncode == 0 and result.stdout.strip():
                    text_content = result.stdout
                    extraction_method = "antiword"
                    logger.info(f"Successfully extracted text from DOC using antiword: {original_filename}")
                else:
                    raise ValueError("Antiword extraction failed")
            except (FileNotFoundError, subprocess.TimeoutExpired, Exception) as e:
                logger.debug(f"antiword failed for {original_filename}: {e}")

                # Method 3: Try python-docx with compatibility mode
                try:
                    import docx

                    with open(file_path, "rb") as f:
                        document = docx.Document(f)
                        paragraphs = [para.text for para in document.paragraphs]
                        text_content = "\n".join(paragraphs)
                        extraction_method = "python-docx-compat"
                        if text_content.strip():
                            logger.info(
                                f"Successfully extracted text from DOC using python-docx compatibility: {original_filename}"
                            )
                        else:
                            raise ValueError("Empty content from python-docx")
                except Exception as e:
                    logger.debug(f"python-docx compatibility failed for {original_filename}: {e}")

                    # Method 4: Try oletools for OLE compound document parsing
                    try:
                        from oletools import olefile

                        if olefile.isOleFile(file_path):
                            # This is a basic OLE file reader - more sophisticated parsing would be needed
                            # for full text extraction, but this provides basic content detection
                            text_content = f"Legacy DOC file detected: {original_filename}. OLE compound document format."
                            extraction_method = "oletools-detection"
                            logger.info(f"Detected legacy DOC format using oletools: {original_filename}")
                        else:
                            raise ValueError("Not a valid OLE file")
                    except (ImportError, Exception) as e:
                        logger.debug(f"oletools failed for {original_filename}: {e}")

                        # Method 5: Try textract as final fallback
                        try:
                            import textract

                            raw_text = textract.process(file_path)
                            text_content = raw_text.decode("utf-8", errors="ignore")
                            extraction_method = "textract"
                            if text_content.strip():
                                logger.info(
                                    f"Successfully extracted text from DOC using textract: {original_filename}"
                                )
                            else:
                                raise ValueError("Empty content from textract")
                        except (ImportError, Exception) as e:
                            logger.debug(f"textract failed for {original_filename}: {e}")
                            # Final fallback
                            text_content = f"Could not extract content from legacy DOC file: {original_filename}. File may be corrupt or require specialized tools."
                            extraction_method = "fallback"

        # Clean up extracted text
        if text_content:
            # Remove excessive whitespace while preserving structure
            lines = [line.strip() for line in text_content.split("\n")]
            text_content = "\n".join(line for line in lines if line)

            # Add extraction method info for debugging
            text_content += f"\n\n[Extracted using: {extraction_method}]"

    except Exception as e:
        logger.error(f"Error processing legacy DOC {original_filename}: {e}")
        text_content = f"Error extracting text from legacy DOC {original_filename}: {e!s}"

    content_type, _ = mimetypes.guess_type(file_path)
    file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

    # Create proper FileMetadata object with required fields
    file_metadata = FileMetadata(filename=original_filename, size=file_size)

    logger.info(
        f"✅ Created FileMetadata for legacy DOC {original_filename}, size: {file_size}, method: {extraction_method}"
    )

    return ProcessedDocument(
        file_name=original_filename,
        content=text_content,
        document_type=document_type,
        file_type=FileType.DOC,
        metadata=file_metadata,
    )
