from __future__ import annotations

import io
import mimetypes

import docx

from utils.data_models import DocumentType, FileType, ProcessedDocument
from utils.logging_config import setup_logging


logger = setup_logging("docx_processor")


async def process_docx(
    file_path: str, document_type: DocumentType, original_filename: str
) -> ProcessedDocument:
    """
    Processes a DOCX file by extracting its content from a given path.
    """
    logger.debug(f"Processing DOCX: {original_filename}")

    try:
        with open(file_path, "rb") as f:
            document = docx.Document(io.BytesIO(f.read()))
            full_text = "\n".join([para.text for para in document.paragraphs])
    except Exception as e:
        logger.info(f"Could not process {original_filename} with python-docx: {e}")
        full_text = f"Could not extract content from {original_filename}."

    content_type, _ = mimetypes.guess_type(file_path)

    return ProcessedDocument(
        file_name=original_filename,
        content=full_text,
        document_type=document_type,
        file_type=FileType.DOCX,
        metadata={"content_type": content_type},
    )
